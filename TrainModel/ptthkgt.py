from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

doc = Document()
def H1(t): doc.add_heading(t, level=1)
def H2(t): doc.add_heading(t, level=2)
def H3(t): doc.add_heading(t, level=3)

doc.add_paragraph("Comparative Study of Algorithms for the Traveling Salesman Problem (TSP)").style='Title'
doc.add_paragraph(f"Date: {datetime.now():%Y-%m-%d}").alignment = WD_ALIGN_PARAGRAPH.CENTER
# (Paste the report sections from above here — you can copy block-by-block.)
# Tip: each section is just doc.add_heading(...) and doc.add_paragraph(...)
doc.save("TSP_Report_Comparative_Study.docx")
